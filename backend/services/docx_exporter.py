"""Markdown-to-DOCX export with Node and python-docx engines."""

from __future__ import annotations

import asyncio
import json
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

try:
    from config import RUNTIME_NODE, TOOLS_DIR
except ModuleNotFoundError:  # Package import used by tests and library consumers.
    from backend.config import RUNTIME_NODE, TOOLS_DIR


class DocxExportError(RuntimeError):
    pass


DEFAULT_SOURCE_BY_TEMPLATE = {
    "thesis_proposal": "PROPOSAL.md",
    "literature_review": "LITERATURE_REVIEW.md",
    "course_paper": "COURSE_PAPER.md",
    "course_report": "COURSE_REPORT.md",
}

DEFAULT_PROFILE_BY_TEMPLATE = {
    "thesis_proposal": "thesis_proposal.json",
    "literature_review": "literature_review.json",
    "course_paper": "course_paper.json",
    "course_report": "course_report.json",
}


def _safe_path(root: Path, supplied: str | Path) -> Path:
    path = Path(supplied)
    if not path.is_absolute():
        path = root / path
    path = path.resolve()
    try:
        path.relative_to(root.resolve())
    except ValueError as exc:
        raise DocxExportError("Path leaves the workflow workspace") from exc
    return path


def resolve_markdown_source(
    workspace: Path, source_file: str | None = None, template: str = ""
) -> Path:
    workspace = workspace.resolve()
    if source_file:
        source = _safe_path(workspace, source_file)
        if source.suffix.lower() not in {".md", ".markdown"}:
            raise DocxExportError("DOCX source must be a Markdown file")
        if not source.is_file():
            raise DocxExportError(f"Markdown source not found: {source_file}")
        return source

    preferred = DEFAULT_SOURCE_BY_TEMPLATE.get(template, "paper/main.md")
    candidates = [
        preferred,
        "paper/main.md",
        "PROPOSAL.md",
        "LITERATURE_REVIEW.md",
        "COURSE_PAPER.md",
        "COURSE_REPORT.md",
    ]
    for relative in dict.fromkeys(candidates):
        path = workspace / relative
        if path.is_file():
            return path
    markdown_files = sorted(
        path for path in workspace.rglob("*.md") if path.is_file() and ".git" not in path.parts
    )
    if markdown_files:
        return markdown_files[0]
    raise DocxExportError("No Markdown source file was found in the workspace")


def resolve_style_profile(style_profile: str | None, template: str = "") -> Path | None:
    profiles = (Path(TOOLS_DIR) / "docx_style_profiles").resolve()
    filename = style_profile or DEFAULT_PROFILE_BY_TEMPLATE.get(
        template, "default_cn_thesis.json"
    )
    candidate = (profiles / filename).resolve()
    try:
        candidate.relative_to(profiles)
    except ValueError as exc:
        raise DocxExportError("Style profile leaves the profile directory") from exc
    if not candidate.is_file():
        if style_profile:
            raise DocxExportError(f"Style profile not found: {style_profile}")
        return None
    return candidate


def _load_profile(path: Path | None) -> dict[str, Any]:
    if path is None:
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise DocxExportError(f"Invalid style profile: {path.name}") from exc
    if not isinstance(value, dict):
        raise DocxExportError(f"Style profile must contain a JSON object: {path.name}")
    return value


def _strip_inline(value: str) -> str:
    value = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\*\*(.*?)\*\*|__(.*?)__", lambda m: m.group(1) or m.group(2), value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)|_([^_]+)_", lambda m: m.group(1) or m.group(2), value)
    return value.replace("`", "").strip()


def _table_cells(line: str) -> list[str]:
    return [_strip_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = line.strip().strip("|").split("|")
    return bool(cells) and all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell) for cell in cells)


def _set_east_asia_font(run: Any, font_name: str) -> None:
    try:
        from docx.oxml.ns import qn

        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except (AttributeError, TypeError):
        pass


def export_markdown_python(
    source: str | Path,
    output: str | Path,
    workspace: str | Path | None = None,
    profile: str | Path | None = None,
) -> dict[str, Any]:
    """Create a valid academic DOCX using only python-docx."""

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise DocxExportError("python-docx is unavailable") from exc

    source_path = Path(source).resolve()
    workspace_path = Path(workspace or source_path.parent).resolve()
    output_path = Path(output).resolve()
    _safe_path(workspace_path, source_path)
    _safe_path(workspace_path, output_path)
    if not source_path.is_file():
        raise DocxExportError(f"Markdown source not found: {source_path}")
    profile_path = Path(profile).resolve() if profile else None
    style = _load_profile(profile_path)
    page = style.get("page") if isinstance(style.get("page"), dict) else {}
    fonts = style.get("fonts") if isinstance(style.get("fonts"), dict) else {}
    body_style = style.get("body") if isinstance(style.get("body"), dict) else {}
    headings = style.get("headings") if isinstance(style.get("headings"), dict) else {}
    image_style = style.get("image") if isinstance(style.get("image"), dict) else {}

    document = Document()
    document.core_properties.author = "ArHub"
    document.core_properties.title = source_path.stem
    section = document.sections[0]
    section.top_margin = Cm(float(page.get("margin_top_cm", 2.5)))
    section.bottom_margin = Cm(float(page.get("margin_bottom_cm", 2.5)))
    section.left_margin = Cm(float(page.get("margin_left_cm", 2.5)))
    section.right_margin = Cm(float(page.get("margin_right_cm", 2.5)))

    body_font = str(fonts.get("chinese_body", "SimSun"))
    heading_font = str(fonts.get("chinese_heading", "SimHei"))
    normal = document.styles["Normal"]
    normal.font.name = str(fonts.get("latin", "Times New Roman"))
    normal.font.size = Pt(float(body_style.get("font_size_pt", 12)))
    normal.paragraph_format.line_spacing = float(body_style.get("line_spacing", 1.5))
    normal.paragraph_format.first_line_indent = Pt(
        float(body_style.get("first_line_indent_chars", 2))
        * float(body_style.get("font_size_pt", 12))
    )
    for level in range(1, 4):
        heading = document.styles[f"Heading {level}"]
        heading.font.name = heading_font
        heading.font.size = Pt(float(headings.get(f"level{level}_pt", 18 - level * 2)))
        heading.font.bold = bool(headings.get("bold", True))

    content = source_path.read_text(encoding="utf-8-sig", errors="replace")
    content = re.sub(r"\A---\s*\n[\s\S]*?\n---\s*\n", "", content, count=1)
    lines = content.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Pt(0)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = str(fonts.get("monospace", "Consolas"))
                run.font.size = Pt(9)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and _is_table_separator(lines[index + 1])
        ):
            rows = [_table_cells(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_table_cells(lines[index]))
                index += 1
            columns = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=columns)
            table.style = "Table Grid"
            for row_number, values in enumerate(rows):
                for column, value in enumerate(values):
                    cell = table.cell(row_number, column)
                    cell.text = value
                    if row_number == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(3, len(heading_match.group(1)))
            paragraph = document.add_heading(_strip_inline(heading_match.group(2)), level=level)
            alignment = str(headings.get(f"level{level}_alignment", "left"))
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if alignment == "center" else WD_ALIGN_PARAGRAPH.LEFT
            )
            for run in paragraph.runs:
                _set_east_asia_font(run, heading_font)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            raw_path = image_match.group(2).split(maxsplit=1)[0].strip("<>")
            try:
                image = _safe_path(workspace_path, source_path.parent / raw_path)
            except DocxExportError:
                image = Path()
            if image.is_file():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.add_run().add_picture(
                    str(image), width=Cm(float(image_style.get("max_width_cm", 14)))
                )
                if image_match.group(1):
                    caption = document.add_paragraph(_strip_inline(image_match.group(1)))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.first_line_indent = Pt(0)
            else:
                document.add_paragraph(_strip_inline(image_match.group(1) or raw_path))
            index += 1
            continue
        bullet = re.match(r"^[-+*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet or numbered:
            paragraph = document.add_paragraph(
                _strip_inline((bullet or numbered).group(1)),
                style="List Bullet" if bullet else "List Number",
            )
            paragraph.paragraph_format.first_line_indent = Pt(0)
            index += 1
            continue
        if re.fullmatch(r"(?:---+|\*\*\*+)", stripped):
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            following = lines[index].strip()
            if (
                not following
                or following.startswith(("#", "```", "|", "!["))
                or re.match(r"^(?:[-+*]\s+|\d+[.)]\s+)", following)
            ):
                break
            paragraph_lines.append(following)
            index += 1
        paragraph = document.add_paragraph(_strip_inline(" ".join(paragraph_lines)))
        for run in paragraph.runs:
            _set_east_asia_font(run, body_font)

    if code_lines:
        document.add_paragraph("\n".join(code_lines))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return {
        "success": True,
        "engine": "python",
        "source_path": source_path,
        "output_path": output_path,
        "size": output_path.stat().st_size,
        "log": f"DOCX exported with python-docx: {output_path.name}",
    }


def _node_executable() -> Path | None:
    candidates = [Path(RUNTIME_NODE) / "node.exe", Path(RUNTIME_NODE), Path("node")]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
        found = shutil.which(str(candidate))
        if found:
            return Path(found)
    return None


def _node_module_paths(converter: Path) -> list[Path]:
    candidates = [converter.parent / "node_modules", Path(RUNTIME_NODE) / "node_modules"]
    runtime = Path(RUNTIME_NODE).resolve()
    if runtime.name.lower() == "node" and len(runtime.parents) >= 2:
        candidates.append(
            runtime.parents[1]
            / "resources"
            / "app"
            / "tools"
            / "docx-cn-engine"
            / "node_modules"
        )
    return [path for path in candidates if path.is_dir()]


async def _run_node(
    node: Path,
    converter: Path,
    source: Path,
    output: Path,
    workspace: Path,
    profile: Path | None,
) -> tuple[int, str]:
    arguments = [
        str(node),
        str(converter),
        "--source",
        str(source),
        "--output",
        str(output),
        "--workspace",
        str(workspace),
    ]
    if profile:
        arguments.extend(["--profile", str(profile)])
    env = {key: str(value) for key, value in os.environ.items() if value is not None}
    module_paths = _node_module_paths(converter)
    if module_paths:
        existing = env.get("NODE_PATH", "")
        env["NODE_PATH"] = os.pathsep.join([*(str(path) for path in module_paths), existing]).rstrip(
            os.pathsep
        )
    kwargs: dict[str, Any] = {}
    if os.name == "nt":
        kwargs["creationflags"] = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    process = await asyncio.create_subprocess_exec(
        *arguments,
        cwd=str(workspace),
        env=env,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
        **kwargs,
    )
    try:
        stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=300)
    except asyncio.TimeoutError:
        process.kill()
        stdout, stderr = await process.communicate()
        return 124, (stdout + stderr).decode("utf-8", errors="replace") + "\nNode export timed out"
    return int(process.returncode or 0), (stdout + stderr).decode("utf-8", errors="replace")


async def export_markdown_to_docx(
    workspace: str | Path,
    *,
    source_file: str | None = None,
    style_profile: str | None = None,
    engine: str | None = "auto",
    template: str = "",
) -> dict[str, Any]:
    workspace_path = Path(workspace).resolve()
    source = resolve_markdown_source(workspace_path, source_file, template)
    profile = resolve_style_profile(style_profile, template)
    output = source.with_suffix(".docx")
    selected = str(engine or "auto").lower()
    if selected not in {"auto", "node", "python"}:
        raise DocxExportError("engine must be auto, node, or python")

    node_error = ""
    if selected in {"auto", "node"}:
        node = _node_executable()
        converter = Path(TOOLS_DIR) / "docx-cn-engine" / "md_to_docx.js"
        if node and converter.is_file():
            output.unlink(missing_ok=True)
            code, node_log = await _run_node(
                node, converter, source, output, workspace_path, profile
            )
            if code == 0 and output.is_file():
                return {
                    "success": True,
                    "engine": "node",
                    "source_path": source,
                    "output_path": output,
                    "size": output.stat().st_size,
                    "log": node_log[-100_000:],
                }
            node_error = node_log[-20_000:] or f"Node engine exited with {code}"
        else:
            node_error = "Node runtime or md_to_docx.js is unavailable"
        if selected == "node":
            raise DocxExportError(node_error)

    result = await asyncio.to_thread(
        export_markdown_python, source, output, workspace_path, profile
    )
    if node_error:
        result["log"] = f"Node engine unavailable; used python fallback.\n{node_error}\n{result['log']}"
    return result
